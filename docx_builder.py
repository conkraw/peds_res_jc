"""Word export logic for Journal Club Builder.

This module creates:
- a compact session summary DOCX
- a mentor review DOCX with full slide text and Track Changes enabled

The visual style intentionally mirrors the printable planning worksheet: blue
section headers, gray field labels, white content areas, and simple table-grid
structure. The goal is readability and standardization without changing the
underlying export content.
"""

from __future__ import annotations

from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile
from typing import Any, Dict, Iterable, List

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
import qrcode

from feedback_config import REDCAP_DISPLAY_URL, REDCAP_QR_URL
from slide_schema import SLIDES


# Colors match the printable planning worksheet style.
BLUE = "4F8ED9"
BLUE_DARK = RGBColor(30, 80, 130)
HEADER_GRAY = "D9D9D9"
FOOTER_BLUE = "DDEFF4"
WHITE = "FFFFFF"
BORDER = "000000"
TEXT_DARK = RGBColor(20, 20, 20)
TEXT_MUTED = RGBColor(95, 95, 95)

EMU_PER_INCH = 914400
TWIPS_PER_INCH = 1440


def _body_width_inches(doc: Document) -> float:
    """Return the usable page width for the current section, in inches."""
    section = doc.sections[-1]
    return float(section.page_width - section.left_margin - section.right_margin) / EMU_PER_INCH


def _set_table_fixed_width(table, width_inches: float) -> None:
    """Lock table width so section banners and content tables line up in Word."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width_inches * TWIPS_PER_INCH)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def _set_cell_width(cell, width_inches: float) -> None:
    """Set cell width at both python-docx and OOXML levels."""
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * TWIPS_PER_INCH)))
    tc_w.set(qn("w:type"), "dxa")


def _lock_table_widths(table, widths: List[float]) -> None:
    """Apply a fixed table layout and fixed column widths to every row."""
    total_width = sum(widths)
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_fixed_width(table, total_width)
    for idx, width in enumerate(widths):
        if idx < len(table.columns):
            table.columns[idx].width = Inches(width)
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                _set_cell_width(row.cells[idx], width)



# -----------------------------
# Basic text helpers
# -----------------------------


def _safe_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _first_line(value: Any) -> str:
    lines = [line.strip() for line in _safe_text(value).splitlines() if line.strip()]
    return lines[0] if lines else ""


def _lines(value: Any, limit: int | None = None) -> List[str]:
    lines = [line.strip() for line in _safe_text(value).splitlines() if line.strip()]
    return lines[:limit] if limit else lines


def _bullet_text(items: Iterable[str], limit: int | None = None) -> str:
    cleaned = [str(item).strip() for item in items if str(item).strip()]
    if limit is not None:
        cleaned = cleaned[:limit]
    return "\n".join(f"• {item}" for item in cleaned)


def _numbered_text(items: Iterable[str], limit: int | None = None) -> str:
    cleaned = [str(item).strip() for item in items if str(item).strip()]
    if limit is not None:
        cleaned = cleaned[:limit]
    return "\n".join(f"{idx}. {item}" for idx, item in enumerate(cleaned, start=1))


# -----------------------------
# DOCX table styling helpers
# -----------------------------


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_borders(cell, color: str = BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _set_cell_margins(cell, top: int = 60, start: int = 80, bottom: int = 60, end: int = 80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, space_after: float = 0, line_spacing: float = 1.0) -> None:
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = line_spacing


def _clear_cell(cell) -> None:
    cell.text = ""
    # Ensure one empty paragraph remains.
    if not cell.paragraphs:
        cell.add_paragraph()


def _write_cell_text(
    cell,
    text: Any,
    *,
    font_size: float = 9.0,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor = TEXT_DARK,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    line_spacing: float = 1.0,
) -> None:
    """Write text into a table cell while preserving line breaks."""
    _clear_cell(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraphs = _safe_text(text).splitlines() or [""]
    first = True
    for raw in paragraphs:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        _format_paragraph(p, align=align, space_after=0, line_spacing=line_spacing)
        run = p.add_run(raw)
        run.font.name = "Arial"
        run.font.size = Pt(font_size)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = color


def _style_table_grid(table) -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for cell in row.cells:
            _set_cell_borders(cell)
            _set_cell_margins(cell)


def _set_table_widths(table, widths: List[float]) -> None:
    """Lock table and column widths in inches."""
    _lock_table_widths(table, widths)


def _pop_pending_banner(doc: Document) -> str:
    """Return and clear a pending section banner.

    Word can show a small visual gap between two separate consecutive tables
    even when paragraph spacing is zero. To make the blue section header touch
    the first content box, _add_banner stores the text and the next table helper
    inserts it as the first row of the same table.
    """
    banner = str(getattr(doc, "_jc_pending_banner", "") or "").strip()
    setattr(doc, "_jc_pending_banner", "")
    return banner


def _add_banner(doc: Document, text: str) -> None:
    """Queue a blue section header to be inserted into the next table.

    This avoids the small gap Microsoft Word can display between separate
    consecutive tables. The following helper (_add_field_box,
    _add_two_column_value_table, etc.) will consume this value and place the
    banner as the first row of the same table as the content.
    """
    setattr(doc, "_jc_pending_banner", _safe_text(text))

def _add_spacer(doc: Document, pts: float = 4) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)
    p.paragraph_format.line_spacing = 1.0


def _add_banner(doc: Document, text: str) -> None:
    """Blue full-width section header, matching the planning worksheet.

    Do not add a spacer after the banner; the following field/table should
    sit directly underneath it so the section looks like one connected block.
    """
    table = doc.add_table(rows=1, cols=1)
    _style_table_grid(table)
    _set_table_widths(table, [_body_width_inches(doc)])
    cell = table.cell(0, 0)
    _shade_cell(cell, BLUE)
    _write_cell_text(
        cell,
        _safe_text(text).upper(),
        font_size=10.5,
        bold=True,
        color=RGBColor(255, 255, 255),
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )


def _add_document_title_block(doc: Document, title: str, subtitle: str = "", kicker: str = "JOURNAL CLUB BUILDER") -> None:
    """Top title table styled like the worksheet header."""
    table = doc.add_table(rows=0, cols=1)
    _style_table_grid(table)

    cell = table.add_row().cells[0]
    _shade_cell(cell, "C8D8EA")
    _write_cell_text(cell, kicker, font_size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    cell = table.add_row().cells[0]
    _shade_cell(cell, HEADER_GRAY)
    _write_cell_text(cell, title, font_size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    if subtitle:
        cell = table.add_row().cells[0]
        _shade_cell(cell, WHITE)
        _write_cell_text(cell, subtitle, font_size=9.5, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    _set_table_widths(table, [_body_width_inches(doc)])
    _add_spacer(doc, 8)


def _add_field_box(
    doc: Document,
    label: str,
    text: Any,
    *,
    footer: str = "",
    content_font_size: float = 9.2,
    content_bold: bool = False,
    content_italic: bool = False,
    label_font_size: float = 9.5,
    keep_blank: bool = True,
) -> None:
    """Planning-form style field: gray label row, white content row, optional pale-blue footer."""
    banner = _pop_pending_banner(doc)
    rows = (1 if banner else 0) + (3 if footer else 2)
    table = doc.add_table(rows=rows, cols=1)
    _style_table_grid(table)

    row_idx = 0
    if banner:
        _add_banner_row(table, banner)
        row_idx = 1

    label_cell = table.cell(row_idx, 0)
    _shade_cell(label_cell, HEADER_GRAY)
    _write_cell_text(label_cell, _safe_text(label).upper(), font_size=label_font_size, bold=True)

    content_cell = table.cell(row_idx + 1, 0)
    _shade_cell(content_cell, WHITE)
    content = _safe_text(text)
    if not content and keep_blank:
        content = "[blank]"
    _write_cell_text(
        content_cell,
        content,
        font_size=content_font_size,
        bold=content_bold,
        italic=content_italic,
        line_spacing=1.05,
    )

    if footer:
        footer_cell = table.cell(row_idx + 2, 0)
        _shade_cell(footer_cell, FOOTER_BLUE)
        _write_cell_text(footer_cell, _safe_text(footer).upper(), font_size=8.2, bold=True)

    _set_table_widths(table, [_body_width_inches(doc)])
    _add_spacer(doc, 5)

def _add_two_column_value_table(doc: Document, rows: List[tuple[str, str]], *, label_width: float = 1.65, value_width: float | None = None) -> None:
    banner = _pop_pending_banner(doc)
    table = doc.add_table(rows=1 if banner else 0, cols=2)
    _style_table_grid(table)
    total_width = _body_width_inches(doc)
    if value_width is None:
        value_width = max(1.0, total_width - label_width)

    if banner:
        _add_banner_row(table, banner)

    for label, value in rows:
        cells = table.add_row().cells
        _shade_cell(cells[0], HEADER_GRAY)
        _shade_cell(cells[1], WHITE)
        _write_cell_text(cells[0], _safe_text(label).upper(), font_size=8.5, bold=True)
        _write_cell_text(cells[1], _safe_text(value), font_size=8.6)
    _set_table_widths(table, [label_width, value_width])
    _add_spacer(doc, 5)

def _add_two_column_text_boxes(doc: Document, left_label: str, left_text: str, right_label: str, right_text: str) -> None:
    banner = _pop_pending_banner(doc)
    table = doc.add_table(rows=(3 if banner else 2), cols=2)
    _style_table_grid(table)

    row_offset = 0
    if banner:
        _add_banner_row(table, banner)
        row_offset = 1

    headers = table.rows[row_offset].cells
    bodies = table.rows[row_offset + 1].cells

    for cell, label in zip(headers, [left_label, right_label]):
        _shade_cell(cell, HEADER_GRAY)
        _write_cell_text(cell, label.upper(), font_size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for cell, text in zip(bodies, [left_text, right_text]):
        _shade_cell(cell, WHITE)
        _write_cell_text(cell, text or "[blank]", font_size=8.5, line_spacing=1.0)

    total_width = _body_width_inches(doc)
    _set_table_widths(table, [total_width / 2, total_width / 2])
    _add_spacer(doc, 5)

def _add_editable_table(doc: Document, label: str, rows: Any) -> None:
    """Add a readable editable representation of a slide table."""
    _add_banner(doc, label)
    if not isinstance(rows, list) or not rows:
        _add_field_box(doc, "Table", "[blank table]", content_italic=True)
        return

    columns: List[str] = []
    for row in rows:
        if isinstance(row, dict):
            for key in row.keys():
                if str(key) not in columns:
                    columns.append(str(key))
    if not columns:
        _add_field_box(doc, "Table", "[blank table]", content_italic=True)
        return

    banner = _pop_pending_banner(doc)
    table = doc.add_table(rows=(2 if banner else 1), cols=len(columns))
    _style_table_grid(table)

    header_row_idx = 0
    if banner:
        _add_banner_row(table, banner)
        header_row_idx = 1

    for idx, column in enumerate(columns):
        cell = table.rows[header_row_idx].cells[idx]
        _shade_cell(cell, HEADER_GRAY)
        _write_cell_text(cell, column.upper(), font_size=8.2, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row in rows:
        if not isinstance(row, dict):
            continue
        cells = table.add_row().cells
        for idx, column in enumerate(columns):
            _shade_cell(cells[idx], WHITE)
            _write_cell_text(cells[idx], _safe_text(row.get(column, "")), font_size=8.2)
    col_width = _body_width_inches(doc) / max(1, len(columns))
    _set_table_widths(table, [col_width] * len(columns))
    _add_spacer(doc, 5)

# -----------------------------
# QR / feedback helpers
# -----------------------------


def _make_qr_image(url: str) -> BytesIO:
    qr = qrcode.QRCode(
        version=None,
        error_correction=qrcode.constants.ERROR_CORRECT_M,
        box_size=8,
        border=2,
    )
    qr.add_data(url)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    bio = BytesIO()
    img.save(bio, format="PNG")
    bio.seek(0)
    return bio


def _add_feedback_block(doc: Document) -> None:
    """Add fixed feedback link and QR code to the one-page summary."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run("Feedback: ")
    run.bold = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = BLUE_DARK
    link = p.add_run(REDCAP_DISPLAY_URL)
    link.font.size = Pt(8.5)
    link.font.color.rgb = BLUE_DARK

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    run = p2.add_run()
    run.add_picture(_make_qr_image(REDCAP_QR_URL), width=Inches(0.72))


# -----------------------------
# One-page / audience summary export
# -----------------------------


def build_word_summary(deck: Dict[str, Dict[str, Any]]) -> BytesIO:
    """Build a compact DOCX summary for session handout/records.

    The included content matches the prior summary export; only the visual
    formatting has been changed to planning-form style tables.
    """
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.38)
    section.bottom_margin = Inches(0.38)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(8.8)

    title_data = deck.get("title_goal", {})
    pico = deck.get("pico", {})
    design = deck.get("study_design", {})
    result = deck.get("main_result", {})
    bottom = deck.get("clinical_bottom_line", {})
    final = deck.get("final_bottom_line", {})

    session_title = _safe_text(title_data.get("session_title", "Journal Club")) or "Journal Club"
    article_title = _safe_text(title_data.get("article_title", ""))
    _add_document_title_block(doc, session_title, article_title, kicker="PEDIATRIC RESIDENCY JOURNAL CLUB BUILDER")

    _add_banner(doc, "Session Purpose")
    _add_field_box(doc, "Teaching Goal", title_data.get("teaching_goal", ""), content_font_size=8.8)

    _add_banner(doc, "Article In One View")
    _add_two_column_value_table(
        doc,
        [
            ("Patient/Problem", _safe_text(pico.get("patient", ""))),
            ("Study Question", _safe_text(pico.get("plain_question", ""))),
            ("Study Design", _safe_text(design.get("design", ""))),
            ("Primary Outcome", _safe_text(pico.get("outcome", ""))),
        ],
    )

    _add_banner(doc, "Main Result")
    _add_field_box(doc, "Headline", result.get("main_result", ""), content_font_size=8.8, content_bold=True)
    _add_field_box(doc, "Plain Language", result.get("plain_result", ""), content_font_size=8.8)

    _add_banner(doc, "Clinical Bottom Line")
    _add_field_box(doc, "Clinical Bottom Line", bottom.get("bottom_line", ""), content_font_size=8.8)
    _add_field_box(doc, "Practice Implication", bottom.get("practice_statement", ""), content_font_size=8.8)

    _add_banner(doc, "Why Trust It / Why Be Cautious")
    trust_text = _bullet_text(_lines(bottom.get("trust_bullets", ""), limit=3))
    caution_text = _bullet_text(_lines(bottom.get("caution_bullets", ""), limit=3))
    _add_two_column_text_boxes(doc, "Trust Factors", trust_text, "Cautions", caution_text)

    _add_banner(doc, "Discussion Questions")
    questions = [
        _safe_text(deck.get("patient_problem", {}).get("discussion_question", "")),
        _safe_text(pico.get("discussion_question", "")),
        _safe_text(result.get("discussion_question", "")),
        _safe_text(deck.get("apply_back", {}).get("return_question", "")),
    ]
    _add_field_box(doc, "Questions For Discussion", _bullet_text([q for q in questions if q], limit=4), content_font_size=8.8)

    _add_banner(doc, "Resident Take-Home")
    _add_field_box(doc, "Take-Home Sentence", final.get("resident_take_home", ""), content_font_size=8.8)

    # _add_feedback_block(doc)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


# -----------------------------
# Mentor review text export
# -----------------------------


def _add_reviewer_guidelines(doc: Document) -> None:
    """Add mentor/reviewer instructions at the top of the review document."""
    _add_banner(doc, "Reviewer Guidelines")
    guidelines = [
        "Use Track Changes and/or comments so the resident can see your suggestions.",
        "Focus on clarity, accuracy, educational value, clinical reasoning, and whether the message is easy to present aloud.",
        "Avoid stylistic preference edits unless they improve readability, reduce confusion, or make the teaching point clearer.",
        "Preserve the resident's voice when possible; suggest targeted edits rather than rewriting the entire slide.",
        "Flag overstatements, missing limitations, unclear applicability, jargon, or places where the clinical takeaway is too broad.",
        "Keep slide text concise. If content is correct but too dense, suggest what to cut or move to facilitator notes.",
    ]
    _add_field_box(doc, "Reviewer Focus", _bullet_text(guidelines), content_font_size=9.0)
    _add_field_box(
        doc,
        "Reviewer Workflow",
        "Edit the text below directly, or add comments beside sections that need discussion. The goal is not to perfect the slide design; the goal is to help the resident make the content clearer, more accurate, and more useful for learners.",
        content_font_size=9.0,
        content_italic=True,
    )


def _add_review_text_block(doc: Document, label: str, text: Any) -> None:
    """Add one editable slide-text field with a clear label."""
    content = _safe_text(text) or "[blank]"
    _add_field_box(doc, label, content, content_font_size=9.3)


def _add_review_table_block(doc: Document, label: str, rows: Any) -> None:
    """Add a readable editable representation of a slide table."""
    if not isinstance(rows, list) or not rows:
        _add_field_box(doc, label, "[blank table]", content_font_size=9.3, content_italic=True)
        return

    columns: List[str] = []
    for row in rows:
        if isinstance(row, dict):
            for key in row.keys():
                if str(key) not in columns:
                    columns.append(str(key))
    if not columns:
        _add_field_box(doc, label, "[blank table]", content_font_size=9.3, content_italic=True)
        return

    _add_banner(doc, label)
    banner = _pop_pending_banner(doc)
    table = doc.add_table(rows=(2 if banner else 1), cols=len(columns))
    _style_table_grid(table)

    header_row_idx = 0
    if banner:
        _add_banner_row(table, banner)
        header_row_idx = 1

    for idx, column in enumerate(columns):
        cell = table.rows[header_row_idx].cells[idx]
        _shade_cell(cell, HEADER_GRAY)
        _write_cell_text(cell, column.upper(), font_size=8.3, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row in rows:
        if not isinstance(row, dict):
            continue
        cells = table.add_row().cells
        for idx, column in enumerate(columns):
            _shade_cell(cells[idx], WHITE)
            _write_cell_text(cells[idx], _safe_text(row.get(column, "")), font_size=8.3)

    footer_row = table.add_row().cells
    footer_row[0].merge(footer_row[-1])
    _shade_cell(footer_row[0], FOOTER_BLUE)
    _write_cell_text(footer_row[0], "Editable table text for mentor review", font_size=8.0, bold=True)
    col_width = _body_width_inches(doc) / max(1, len(columns))
    _set_table_widths(table, [col_width] * len(columns))
    _add_spacer(doc, 5)

def _enable_track_changes(docx_stream: BytesIO) -> BytesIO:
    """Turn on Word track-revisions setting for the generated review document.

    This does not create tracked changes by itself. It nudges Word to open the
    document with Track Changes enabled so mentor edits are easier for the
    resident to review.
    """
    try:
        source = BytesIO(docx_stream.getvalue())
        target = BytesIO()
        with ZipFile(source, "r") as zin, ZipFile(target, "w", ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/settings.xml":
                    xml = data.decode("utf-8")
                    if "w:trackRevisions" not in xml:
                        xml = xml.replace("</w:settings>", "<w:trackRevisions/></w:settings>")
                    data = xml.encode("utf-8")
                zout.writestr(item, data)
        target.seek(0)
        return target
    except Exception:
        docx_stream.seek(0)
        return docx_stream


def build_review_text_docx(deck: Dict[str, Dict[str, Any]]) -> BytesIO:
    """Build an editable DOCX containing the PowerPoint text for mentor review.

    This is intentionally different from the summary. It includes the full text
    from the slide-builder fields, organized by slide title, so a mentor can
    provide tracked edits or comments before the resident finalizes the
    PowerPoint. The content is unchanged; the styling is planning-form inspired.
    """
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    title_data = deck.get("title_goal", {})
    session_title = _safe_text(title_data.get("session_title", "Journal Club")) or "Journal Club"
    article_title = _safe_text(title_data.get("article_title", ""))

    _add_document_title_block(
        doc,
        "PowerPoint Text Review",
        f"{session_title}\n{article_title}" if article_title else session_title,
        kicker="JOURNAL CLUB BUILDER",
    )

    _add_reviewer_guidelines(doc)

    for slide_number, slide in enumerate(SLIDES, start=1):
        slide_id = slide["id"]
        slide_data = deck.get(slide_id, {}) or {}
        slide_title = str(slide.get("export_title") or slide.get("label") or slide_id).strip()
        _add_banner(doc, f"Slide {slide_number}: {slide_title}")

        for field in slide.get("fields", []):
            # Respect simple show_if conditions so the review file matches the visible app fields.
            show_if = field.get("show_if") or {}
            if show_if:
                visible = all(slide_data.get(key) == expected for key, expected in show_if.items())
                if not visible:
                    continue

            field_label = str(field.get("label") or field.get("key") or "Field")
            field_key = field.get("key")
            value = slide_data.get(field_key, "") if field_key else ""

            if field.get("type") == "table":
                _add_review_table_block(doc, field_label, value)
            else:
                _add_review_text_block(doc, field_label, value)

        _add_field_box(
            doc,
            "Mentor notes / comments",
            "[Add comments here or use Word comments in the margin.]",
            content_font_size=9.0,
            content_italic=True,
        )

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return _enable_track_changes(output)
