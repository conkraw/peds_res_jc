"""Printable planning form for Journal Club Builder.

This creates a pen-and-paper worksheet that mirrors the Streamlit fields. It is
meant for residents who prefer to read the article on paper, write notes by
hand, then transfer the final answers into the app.
"""

from __future__ import annotations

from io import BytesIO
from typing import Any, Dict, List

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from slide_schema import SLIDES, make_default_deck


BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
GRAY = "F2F2F2"
DARK_GRAY = "555555"


def _safe_text(value: Any) -> str:
    return "" if value is None else str(value).strip()


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_width(cell, width_inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _format_run(run, size: float = 9, bold: bool = False, color: str | None = None) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _add_text(paragraph, text: str, size: float = 9, bold: bool = False, color: str | None = None):
    run = paragraph.add_run(text)
    _format_run(run, size=size, bold=bold, color=color)
    return run


def _limits_text(field: Dict[str, Any]) -> str:
    parts: List[str] = []
    if field.get("required"):
        parts.append("Required")
    else:
        parts.append("Optional")

    if "max_words" in field:
        parts.append(f"max {field['max_words']} words")
    if "max_chars" in field:
        parts.append(f"max {field['max_chars']} characters")
    if "max_lines" in field:
        parts.append(f"max {field['max_lines']} lines")
    if "max_words_per_line" in field:
        parts.append(f"max {field['max_words_per_line']} words/line")
    if "max_rows" in field:
        parts.append(f"max {field['max_rows']} rows")
    if field.get("type") == "number":
        if field.get("min_value") is not None:
            parts.append(f"min {field['min_value']}")
        if field.get("max_value") is not None:
            parts.append(f"max {field['max_value']}")
    return " | ".join(parts)


def _writing_lines_for_field(field: Dict[str, Any]) -> int:
    if field.get("type") == "text":
        return 1
    if "max_lines" in field:
        return min(max(int(field["max_lines"]), 1), 8)
    max_words = int(field.get("max_words", 35) or 35)
    if max_words <= 20:
        return 1
    if max_words <= 35:
        return 2
    if max_words <= 55:
        return 3
    if max_words <= 75:
        return 4
    return 5


def _add_section_heading(doc: Document, nav_name: str, slide_title: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    _set_cell_shading(cell, BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    _add_text(p, nav_name, size=11, bold=True, color="FFFFFF")
    if slide_title and slide_title != nav_name:
        _add_text(p, f"  |  {slide_title}", size=9, bold=False, color="FFFFFF")


def _add_field_box(doc: Document, field: Dict[str, Any]) -> None:
    table = doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    header = table.cell(0, 0)
    _set_cell_shading(header, GRAY)
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    _add_text(p, field["label"], size=9.5, bold=True, color=BLUE)
    _add_text(p, f"  ({_limits_text(field)})", size=8, bold=False, color=DARK_GRAY)

    guide = _safe_text(field.get("guide"))
    if guide:
        p2 = header.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        _add_text(p2, guide, size=7.5, color=DARK_GRAY)

    body = table.cell(1, 0)
    body.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    lines = _writing_lines_for_field(field)
    for i in range(lines):
        p = body.paragraphs[0] if i == 0 else body.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        _add_text(p, "_" * 96, size=8, color="777777")

    p = body.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if "max_words" in field:
        _add_text(p, f"Word count: ______ / {field['max_words']} words", size=7.5, color=DARK_GRAY)
    elif "max_lines" in field:
        _add_text(p, f"Line count: ______ / {field['max_lines']} lines", size=7.5, color=DARK_GRAY)
    elif field.get("type") == "number":
        _add_text(p, "Numeric value: ______", size=7.5, color=DARK_GRAY)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def _add_select_field(doc: Document, field: Dict[str, Any]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    _set_cell_shading(cell, GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    _add_text(p, field["label"], size=9.5, bold=True, color=BLUE)
    _add_text(p, f"  ({_limits_text(field)})", size=8, color=DARK_GRAY)

    guide = _safe_text(field.get("guide"))
    if guide:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        _add_text(p2, guide, size=7.5, color=DARK_GRAY)

    p3 = cell.add_paragraph()
    p3.paragraph_format.space_after = Pt(0)
    options = field.get("options", [])
    _add_text(p3, "   ".join(f"[ ] {option}" for option in options), size=8.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def _parse_columns(raw: Any) -> List[str]:
    if isinstance(raw, list):
        columns = raw
    else:
        columns = str(raw or "").replace("\n", ",").split(",")
    out: List[str] = []
    for col in columns:
        cleaned = str(col).strip()
        if cleaned and cleaned not in out:
            out.append(cleaned)
    return out


def _add_table_field(doc: Document, field: Dict[str, Any], slide_data: Dict[str, Any]) -> None:
    max_rows = int(field.get("max_rows", 5) or 5)
    columns_key = field.get("columns_key")
    columns = _parse_columns(slide_data.get(columns_key, "")) if columns_key else []
    if not columns:
        default_rows = field.get("default", [])
        if default_rows and isinstance(default_rows, list) and isinstance(default_rows[0], dict):
            columns = list(default_rows[0].keys())
    if not columns:
        columns = ["Column 1", "Column 2", "Column 3"]

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(2)
    title.paragraph_format.space_after = Pt(1)
    _add_text(title, field["label"], size=9.5, bold=True, color=BLUE)
    _add_text(title, f"  ({_limits_text(field)})", size=8, color=DARK_GRAY)

    guide = _safe_text(field.get("guide"))
    if guide:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        _add_text(p, guide, size=7.5, color=DARK_GRAY)

    table = doc.add_table(rows=max_rows + 1, cols=len(columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    _set_repeat_table_header(table.rows[0])
    for idx, col in enumerate(columns):
        cell = table.cell(0, idx)
        _set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        _add_text(p, col, size=7.5, bold=True, color=BLUE)

    for row_idx in range(1, max_rows + 1):
        for col_idx in range(len(columns)):
            cell = table.cell(row_idx, col_idx)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            _add_text(p, "\n\n", size=7.5)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _add_text(p, f"Rows used: ______ / {max_rows}", size=7.5, color=DARK_GRAY)


def _field_visible(field: Dict[str, Any], slide_data: Dict[str, Any]) -> bool:
    condition = field.get("show_if")
    if not condition:
        return True
    for key, expected in condition.items():
        # The paper form should show conditional fields as options when the current/default
        # selection matches; otherwise it omits the hidden app fields to mirror the app.
        if slide_data.get(key) != expected:
            return False
    return True


def build_printable_planning_form(deck: Dict[str, Dict[str, Any]] | None = None) -> BytesIO:
    """Build a printable DOCX worksheet matching the current app schema."""
    deck = deck or make_default_deck()

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    _add_text(p, "Journal Club Builder", size=17, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    _add_text(p, "Printable Planning Worksheet", size=12, bold=True, color=DARK_GRAY)

    info = doc.add_table(rows=2, cols=4)
    info.style = "Table Grid"
    labels = ["Presenter", "Session title", "Article", "Date"]
    for col_idx, label in enumerate(labels):
        cell = info.cell(0, col_idx)
        _set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        _add_text(p, label, size=8, bold=True, color=BLUE)
        p = info.cell(1, col_idx).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        _add_text(p, "\n", size=8)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    _add_text(
        p,
        "Use this worksheet while reading the article. The field names and limits match the Streamlit app. "
        "After drafting by hand, transfer your final answers into the app and export the PowerPoint.",
        size=8.5,
        color=DARK_GRAY,
    )

    for slide_index, slide in enumerate(SLIDES):
        if slide_index > 0:
            doc.add_page_break()

        slide_data = deck.get(slide["id"], {})
        nav = slide.get("label", slide["id"])
        title = slide.get("export_title", nav)
        _add_section_heading(doc, nav, title)

        for field in slide["fields"]:
            if not _field_visible(field, slide_data):
                continue
            ftype = field.get("type")
            if ftype == "select":
                _add_select_field(doc, field)
            elif ftype == "table":
                _add_table_field(doc, field, slide_data)
            else:
                _add_field_box(doc, field)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output
